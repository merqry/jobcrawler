#!/usr/bin/env python3
"""to_docx.py — Read JSON from stdin and produce a formatted .docx file.

Supported document types:
  - "resume"   : Executive-style resume
  - "outreach" : Multi-message outreach sequence

Usage:
    cat payload.json | python to_docx.py

Exits 0 on success; prints output path to stdout.
Exits non-zero on error; prints message to stderr.
"""

from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

FONT = "Calibri"

# Letter page: 8.5" wide; 0.6" L/R margins → 7.3" text width
PAGE_LEFT_MARGIN = 0.6
PAGE_RIGHT_MARGIN = 0.6
PAGE_TOP_MARGIN = 0.5
PAGE_BOTTOM_MARGIN = 0.5
TEXT_WIDTH = 8.5 - PAGE_LEFT_MARGIN - PAGE_RIGHT_MARGIN  # 7.3"


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------


def _apply_bottom_border(para: Any, sz: str = "6", color: str = "auto") -> None:
    """Add a thin bottom border to a paragraph (renders as a horizontal rule)."""
    pPr = para._p.get_or_add_pPr()
    # Remove any existing pBdr to avoid duplicates
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _remove_default_paragraph(doc: Document) -> None:
    """Remove the single empty paragraph that python-docx inserts by default."""
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)


# ---------------------------------------------------------------------------
# Paragraph / run helpers
# ---------------------------------------------------------------------------


def _set_spacing(
    para: Any,
    before_pt: float = 0,
    after_pt: float = 0,
    line_pt: float | None = None,
) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_pt is not None:
        from docx.enum.text import WD_LINE_SPACING
        pf.line_spacing = Pt(line_pt)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _run(
    para: Any,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size_pt: float | None = None,
    font: str = FONT,
) -> Any:
    """Add a formatted run to *para* and return it."""
    r = para.add_run(text)
    r.font.name = font
    r.font.bold = bold
    r.font.italic = italic
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    return r


def _new_para(
    doc: Document,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: float = 0,
    after_pt: float = 0,
    line_pt: float | None = None,
) -> Any:
    """Add a new paragraph with baseline formatting applied."""
    para = doc.add_paragraph()
    para.alignment = align
    _set_spacing(para, before_pt, after_pt, line_pt)
    return para


def _set_default_style(doc: Document) -> None:
    """Reset the Normal style to Calibri, zero paragraph spacing, multiple 1.1 line spacing."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.1  # Multiple 1.1 — setter auto-applies WD_LINE_SPACING.MULTIPLE


# ---------------------------------------------------------------------------
# Composite helpers
# ---------------------------------------------------------------------------


def _add_section_heading(
    doc: Document,
    text: str,
    size_pt: float = 12,
    before_pt: float = 5,
    after_pt: float = 1,
    rule: bool = True,
) -> Any:
    """Section heading: bold small caps, optional bottom-border rule."""
    para = _new_para(doc, before_pt=before_pt, after_pt=after_pt)
    r = _run(para, text, bold=True, size_pt=size_pt)
    r.font.small_caps = True
    if rule:
        _apply_bottom_border(para, sz="4")
    return para


def _add_role_line(
    doc: Document,
    left_text: str,
    right_text: str,
    left_bold: bool = True,
    left_size: float = 10,
    right_size: float = 10,
    before_pt: float = 3,
    after_pt: float = 0,
) -> Any:
    """
    Render "left_text [TAB] right_text" where the tab stop is right-aligned
    at the full text width, producing a two-column effect on a single line.
    """
    para = _new_para(doc, before_pt=before_pt, after_pt=after_pt)
    para.paragraph_format.tab_stops.add_tab_stop(
        Inches(TEXT_WIDTH), WD_TAB_ALIGNMENT.RIGHT
    )
    _run(para, left_text, bold=left_bold, size_pt=left_size)
    tab = para.add_run("\t")
    tab.font.name = FONT
    tab.font.size = Pt(right_size)
    _run(para, right_text, size_pt=right_size)
    return para


def _add_bullet(
    doc: Document,
    text: str,
    size_pt: float = 10,
    after_pt: float = 1,
) -> Any:
    """Bullet with hanging indent: bullet glyph at left margin, text indented."""
    para = _new_para(doc, after_pt=after_pt)
    pf = para.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.18)
    # Tab stop matches the left indent so text aligns after the bullet
    para.paragraph_format.tab_stops.add_tab_stop(
        Inches(0.18), WD_TAB_ALIGNMENT.LEFT
    )
    _run(para, "\u2022\t" + text, size_pt=size_pt)
    return para


def _add_hrule(doc: Document, before_pt: float = 4, after_pt: float = 4) -> Any:
    """Standalone horizontal rule paragraph."""
    para = _new_para(doc, before_pt=before_pt, after_pt=after_pt)
    _apply_bottom_border(para, sz="6")
    return para


def _add_hyperlink(para: Any, text: str, url: str, size_pt: float = 10) -> None:
    """Append a clickable hyperlink run to *para* (blue, underlined)."""
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    font_elem = OxmlElement("w:rFonts")
    font_elem.set(qn("w:ascii"), FONT)
    font_elem.set(qn("w:hAnsi"), FONT)
    rPr.append(font_elem)

    sz_elem = OxmlElement("w:sz")
    sz_elem.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz_elem)
    szCs_elem = OxmlElement("w:szCs")
    szCs_elem.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szCs_elem)

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), "0563C1")
    rPr.append(color_elem)

    u_elem = OxmlElement("w:u")
    u_elem.set(qn("w:val"), "single")
    rPr.append(u_elem)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Resume builder
# ---------------------------------------------------------------------------


def build_resume(doc: Document, data: dict[str, Any]) -> None:
    _set_default_style(doc)
    _remove_default_paragraph(doc)

    # --- Page margins ---
    section = doc.sections[0]
    section.left_margin = Inches(PAGE_LEFT_MARGIN)
    section.right_margin = Inches(PAGE_RIGHT_MARGIN)
    section.top_margin = Inches(PAGE_TOP_MARGIN)
    section.bottom_margin = Inches(PAGE_BOTTOM_MARGIN)

    # --- Header: Name ---
    hdr = data.get("header", {})
    name_para = _new_para(
        doc, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=1
    )
    _run(name_para, hdr.get("name", ""), bold=True, size_pt=14)

    # --- Header: Contact line ---
    phone = hdr.get("phone", "")
    email = hdr.get("email", "")
    linkedin = hdr.get("linkedin_url", "")
    contact_parts = [p for p in [phone, email, linkedin] if p]
    contact_line = "  |  ".join(contact_parts)

    contact_para = _new_para(
        doc, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=3
    )
    _run(contact_para, contact_line, size_pt=10)
    # Horizontal rule under the header block
    _apply_bottom_border(contact_para, sz="6")

    # --- Summary ---
    summary = data.get("summary", "").strip()
    if summary:
        _add_section_heading(doc, "Summary")
        body = _new_para(doc, before_pt=1, after_pt=1)
        _run(body, summary, size_pt=10)

    # --- Core Strengths ---
    strengths = data.get("core_strengths", [])
    if strengths:
        _add_section_heading(doc, "Core Strengths")
        body = _new_para(doc, before_pt=1, after_pt=1)
        _run(body, "  |  ".join(strengths), size_pt=10)

    # --- Professional Experience ---
    experience = data.get("experience", [])
    if experience:
        _add_section_heading(doc, "Professional Experience")
        for role in experience:
            company = role.get("company", "")
            title = role.get("title", "")
            location = role.get("location", "")
            dates = role.get("dates", "")
            scope = role.get("scope", "").strip()
            sections = role.get("sections", [])

            _add_role_line(
                doc,
                left_text=f"{company} \u2014 {title}",
                right_text=f"{location}  |  {dates}",
                before_pt=4,
            )

            if scope:
                scope_para = _new_para(doc, before_pt=1, after_pt=1)
                _run(scope_para, scope, italic=True, size_pt=10)

            for sec in sections:
                subheading = sec.get("subheading", "").strip()
                bullets = sec.get("bullets", [])
                if subheading:
                    sub_para = _new_para(doc, before_pt=2, after_pt=0)
                    _run(sub_para, subheading, bold=True, italic=True, size_pt=10)
                for bullet in bullets:
                    _add_bullet(doc, bullet, size_pt=10)

    # --- Early Career ---
    early_career = data.get("early_career", [])
    if early_career:
        _add_section_heading(doc, "Early Career \u2014 Product & Engineering")
        for role in early_career:
            company = role.get("company", "")
            title = role.get("title", "")
            location = role.get("location", "")
            dates = role.get("dates", "")
            bullets = role.get("bullets", [])

            _add_role_line(
                doc,
                left_text=f"{company} \u2014 {title}",
                right_text=f"{location}  |  {dates}",
                before_pt=4,
            )
            for bullet in bullets:
                _add_bullet(doc, bullet, size_pt=10)

    # --- Education ---
    education = data.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        for entry in education:
            school = entry.get("school", "")
            degree = entry.get("degree", "")
            location = entry.get("location", "")
            graddate = entry.get("graddate", "")

            right_text = "  |  ".join(p for p in [location, graddate] if p)
            _add_role_line(
                doc,
                left_text=school,
                right_text=right_text,
                before_pt=3,
                after_pt=0,
            )
            if degree:
                degree_para = _new_para(doc, before_pt=0, after_pt=1)
                _run(degree_para, degree, size_pt=10)

    # --- Thought Leadership (extras) ---
    extras = data.get("extras", [])
    if extras:
        _add_section_heading(doc, "Thought Leadership")
        for extra in extras:
            title = extra.get("title", "")
            url = extra.get("url", "").strip()
            description = extra.get("description", "")
            date = extra.get("date", "")

            extra_para = _new_para(doc, before_pt=3, after_pt=1)
            extra_para.paragraph_format.tab_stops.add_tab_stop(
                Inches(TEXT_WIDTH), WD_TAB_ALIGNMENT.RIGHT
            )
            if url:
                _add_hyperlink(extra_para, title, url, size_pt=10)
            else:
                _run(extra_para, title, bold=True, size_pt=10)
            right_parts = [p for p in [description, date] if p]
            if right_parts:
                tab = extra_para.add_run("\t")
                tab.font.name = FONT
                tab.font.size = Pt(9)
                _run(extra_para, "  |  ".join(right_parts), size_pt=10)


# ---------------------------------------------------------------------------
# Outreach builder
# ---------------------------------------------------------------------------


def build_outreach(doc: Document, data: dict[str, Any]) -> None:
    _set_default_style(doc)
    _remove_default_paragraph(doc)

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    contact_name = data.get("contact_name", "")
    slug = data.get("slug", "")

    # Title
    title_para = _new_para(
        doc, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=8
    )
    _run(
        title_para,
        f"Outreach: {contact_name} \u2014 {slug}",
        bold=True,
        size_pt=13,
    )
    _apply_bottom_border(title_para, sz="6")

    messages = data.get("messages", [])
    for i, msg in enumerate(messages):
        label = msg.get("label", f"Message {i + 1}")
        body = msg.get("body", "")

        # Label heading
        label_para = _new_para(doc, before_pt=10 if i > 0 else 6, after_pt=4)
        _run(label_para, label, bold=True, size_pt=12)

        # Body text — preserve line breaks from the source
        for line in body.splitlines():
            line_para = _new_para(doc, before_pt=0, after_pt=3)
            _run(line_para, line, size_pt=11)

        # Separator rule after each message except the last
        if i < len(messages) - 1:
            _add_hrule(doc, before_pt=8, after_pt=0)


# ---------------------------------------------------------------------------
# Output path resolution
# ---------------------------------------------------------------------------


def resolve_output_path(output_dir: str, doc_type: str, slug: str) -> Path:
    base_dir = Path(output_dir).expanduser().resolve()
    base_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{doc_type}_{slug}.docx"
    path = base_dir / filename
    if path.exists():
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{doc_type}_{slug}_{ts}.docx"
        path = base_dir / filename
    return path


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def main() -> int:
    try:
        data: dict[str, Any] = json.load(sys.stdin)
    except json.JSONDecodeError as exc:
        print(f"Error: invalid JSON on stdin: {exc}", file=sys.stderr)
        return 1

    doc_type = data.get("type", "")
    if doc_type not in ("resume", "outreach"):
        print(
            f"Error: 'type' must be 'resume' or 'outreach', got {doc_type!r}",
            file=sys.stderr,
        )
        return 1

    slug = data.get("slug", "document")
    output_dir = data.get("output_dir", ".")

    doc = Document()

    try:
        if doc_type == "resume":
            build_resume(doc, data)
        else:
            build_outreach(doc, data)
    except Exception as exc:  # noqa: BLE001
        print(f"Error building document: {exc}", file=sys.stderr)
        return 1

    out_path = resolve_output_path(output_dir, doc_type, slug)

    try:
        doc.save(str(out_path))
    except OSError as exc:
        print(f"Error saving file: {exc}", file=sys.stderr)
        return 1

    print(str(out_path))
    return 0


if __name__ == "__main__":
    sys.exit(main())
